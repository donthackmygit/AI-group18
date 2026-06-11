from __future__ import annotations

import argparse
import csv
import io
import logging
import os
import re
import sys
import tempfile
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import urlparse

import pandas as pd
import requests
from bs4 import BeautifulSoup, Comment
from docx import Document
from pypdf import PdfReader

logging.getLogger("pypdf").setLevel(logging.ERROR)

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

try:
    import fitz  # type: ignore
except ImportError:
    fitz = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import pytesseract
except ImportError:
    pytesseract = None


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REGISTRY_PATH = PROJECT_ROOT / "data" / "metadata" / "document_registry.xlsx"
OUTPUT_DIR = PROJECT_ROOT / "data" / "processed" / "extracted_text"
LOG_PATH = PROJECT_ROOT / "data" / "processed" / "document_loader_log.csv"

DEFAULT_USER_AGENT = (
    "AI-group18-Tax-RAG-DocumentLoader/1.0 "
    "(educational project; contact: repository owner)"
)

DEFAULT_TIMEOUT_SECONDS = 30
MAX_DOWNLOAD_BYTES = 80 * 1024 * 1024  # 80 MB

DEFAULT_OCR_LANGUAGE = "vie+eng"
DEFAULT_OCR_DPI = 220
DEFAULT_OCR_MIN_CHARS = 80

SUPPORTED_IMAGE_SUFFIXES = {
    ".png",
    ".jpg",
    ".jpeg",
    ".tif",
    ".tiff",
    ".bmp",
    ".webp",
}


@dataclass
class ExtractedDocument:
    text: str
    extractor: str
    page_count: int | None = None
    warning: str = ""


@dataclass
class WebFetchedContent:
    body: bytes
    content_type: str
    final_url: str
    encoding: str | None = None


@dataclass
class LoaderOptions:
    timeout_seconds: int = DEFAULT_TIMEOUT_SECONDS
    user_agent: str = DEFAULT_USER_AGENT
    ocr_enabled: bool = True
    force_ocr: bool = False
    ocr_language: str = DEFAULT_OCR_LANGUAGE
    ocr_dpi: int = DEFAULT_OCR_DPI
    ocr_min_chars: int = DEFAULT_OCR_MIN_CHARS


def is_url(value: str) -> bool:
    parsed = urlparse(value.strip())
    return parsed.scheme in {"http", "https"} and bool(parsed.netloc)


def join_warnings(*warnings: str) -> str:
    return " | ".join(warning for warning in warnings if warning)


def meaningful_text_length(text: str) -> int:
    without_page_markers = re.sub(r"===== PAGE \d+ =====", "", text)
    without_table_markers = re.sub(r"===== TABLE \d+ =====", "", without_page_markers)
    without_errors = re.sub(r"\[PAGE_EXTRACTION_ERROR].*", "", without_table_markers)
    return len(without_errors.strip())


def configure_tesseract_from_env() -> None:
    if pytesseract is None:
        return

    tesseract_cmd = os.getenv("TESSERACT_CMD", "").strip()

    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd


def ensure_ocr_dependencies() -> None:
    if fitz is None:
        raise RuntimeError("OCR for PDF requires PyMuPDF. Install package: pymupdf")

    if Image is None:
        raise RuntimeError("OCR requires Pillow. Install package: pillow")

    if pytesseract is None:
        raise RuntimeError("OCR requires pytesseract. Install package: pytesseract")

    configure_tesseract_from_env()


def extract_text_from_pdf_without_ocr(file_path: Path) -> ExtractedDocument:
    """Extract raw text from a PDF, preserving page boundaries."""
    if fitz is not None:
        text_parts: list[str] = []

        with fitz.open(file_path) as pdf:
            for page_index, page in enumerate(pdf, start=1):
                text_parts.append(f"\n\n===== PAGE {page_index} =====\n")
                text_parts.append(page.get_text("text") or "")

            page_count = len(pdf)

        return ExtractedDocument(
            text="\n".join(text_parts),
            extractor="pymupdf",
            page_count=page_count,
        )

    reader = PdfReader(str(file_path))
    text_parts = []

    for page_index, page in enumerate(reader.pages, start=1):
        text_parts.append(f"\n\n===== PAGE {page_index} =====\n")

        try:
            text_parts.append(page.extract_text() or "")
        except Exception as exc:
            text_parts.append("")
            text_parts.append(f"\n[PAGE_EXTRACTION_ERROR] {exc}\n")

    return ExtractedDocument(
        text="\n".join(text_parts),
        extractor="pypdf",
        page_count=len(reader.pages),
        warning="PyMuPDF is not installed; used pypdf fallback.",
    )


def ocr_pdf_page(
    page,
    page_index: int,
    dpi: int,
    language: str,
) -> str:
    scale = dpi / 72
    matrix = fitz.Matrix(scale, scale)

    pixmap = page.get_pixmap(matrix=matrix, alpha=False)
    image_bytes = pixmap.tobytes("png")

    image = Image.open(io.BytesIO(image_bytes))
    text = pytesseract.image_to_string(
        image,
        lang=language,
        config="--psm 6",
    )

    return f"\n\n===== PAGE {page_index} OCR =====\n{text.strip()}"


def extract_text_from_pdf_with_ocr(
    file_path: Path,
    ocr_language: str = DEFAULT_OCR_LANGUAGE,
    ocr_dpi: int = DEFAULT_OCR_DPI,
) -> ExtractedDocument:
    """OCR all pages of a scanned PDF."""
    ensure_ocr_dependencies()

    text_parts: list[str] = []

    with fitz.open(file_path) as pdf:
        for page_index, page in enumerate(pdf, start=1):
            try:
                text_parts.append(
                    ocr_pdf_page(
                        page=page,
                        page_index=page_index,
                        dpi=ocr_dpi,
                        language=ocr_language,
                    )
                )
            except Exception as exc:
                text_parts.append(f"\n\n===== PAGE {page_index} OCR ERROR =====\n{exc}")

        page_count = len(pdf)

    return ExtractedDocument(
        text="\n".join(text_parts),
        extractor=f"ocr-pdf-pytesseract-{ocr_language}",
        page_count=page_count,
        warning=f"OCR applied with language={ocr_language}, dpi={ocr_dpi}.",
    )


def extract_text_from_pdf(
    file_path: Path,
    options: LoaderOptions,
) -> ExtractedDocument:
    normal_extraction = extract_text_from_pdf_without_ocr(file_path)
    normal_text_length = meaningful_text_length(normal_extraction.text)

    should_run_ocr = (
        options.ocr_enabled
        and (
            options.force_ocr
            or normal_text_length < options.ocr_min_chars
        )
    )

    if not should_run_ocr:
        return normal_extraction

    try:
        ocr_extraction = extract_text_from_pdf_with_ocr(
            file_path=file_path,
            ocr_language=options.ocr_language,
            ocr_dpi=options.ocr_dpi,
        )

        ocr_text_length = meaningful_text_length(ocr_extraction.text)

        if options.force_ocr or ocr_text_length > normal_text_length:
            return ExtractedDocument(
                text=ocr_extraction.text,
                extractor=f"{normal_extraction.extractor}+{ocr_extraction.extractor}",
                page_count=ocr_extraction.page_count,
                warning=join_warnings(
                    normal_extraction.warning,
                    (
                        f"PDF text extraction produced only {normal_text_length} chars; "
                        f"OCR produced {ocr_text_length} chars."
                    ),
                    ocr_extraction.warning,
                ),
            )

        return ExtractedDocument(
            text=normal_extraction.text,
            extractor=normal_extraction.extractor,
            page_count=normal_extraction.page_count,
            warning=join_warnings(
                normal_extraction.warning,
                (
                    f"OCR was attempted but not used because normal extraction "
                    f"had {normal_text_length} chars and OCR had {ocr_text_length} chars."
                ),
            ),
        )

    except Exception as exc:
        return ExtractedDocument(
            text=normal_extraction.text,
            extractor=normal_extraction.extractor,
            page_count=normal_extraction.page_count,
            warning=join_warnings(
                normal_extraction.warning,
                f"OCR failed: {exc}",
            ),
        )


def extract_text_from_image(
    file_path: Path,
    options: LoaderOptions,
) -> ExtractedDocument:
    if not options.ocr_enabled:
        raise ValueError(
            f"Image file requires OCR, but OCR is disabled: {file_path}"
        )

    ensure_ocr_dependencies()

    image = Image.open(file_path)
    text = pytesseract.image_to_string(
        image,
        lang=options.ocr_language,
        config="--psm 6",
    )

    return ExtractedDocument(
        text=text.strip(),
        extractor=f"ocr-image-pytesseract-{options.ocr_language}",
        page_count=1,
        warning=(
            f"Image OCR applied with language={options.ocr_language}, "
            f"dpi=original."
        ),
    )


def extract_text_from_docx(file_path: Path) -> ExtractedDocument:
    """Extract text from a DOCX file, including simple table content."""
    document = Document(file_path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            parts.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"\n===== TABLE {table_index} =====")

        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]

            if any(values):
                parts.append("\t".join(values))

    return ExtractedDocument(
        text="\n".join(parts),
        extractor="python-docx",
        page_count=None,
    )


def _is_useful_doc_line(line: str) -> bool:
    if len(line) < 2:
        return False

    binary_markers = (
        "Root Entry",
        "WordDocument",
        "SummaryInformation",
        "DocumentSummaryInformation",
        "CompObj",
        "Ole10Native",
        "ObjectPool",
        "ObjInfo",
    )

    if any(marker in line for marker in binary_markers):
        return False

    if line.count("ÿ") >= 3 or line.count("\ufffd") >= 1:
        return False

    letters = [char for char in line if char.isalpha()]
    latin_letters = [
        char for char in letters if "LATIN" in unicodedata.name(char, "")
    ]

    digits = sum(1 for char in line if char.isdigit())
    punctuation = sum(1 for char in line if char in ".,;:/()-[]% \"'“”‘’+-=_")
    useful = len(latin_letters) + digits + punctuation

    if letters and len(latin_letters) / len(letters) < 0.8:
        return False

    return useful / max(len(line), 1) >= 0.55 and len(latin_letters) + digits >= 2


def extract_text_from_legacy_doc(file_path: Path) -> ExtractedDocument:
    """Best-effort text extraction for old binary .doc files."""
    raw = file_path.read_bytes()
    decoded = raw.decode("utf-16le", errors="ignore")
    decoded = decoded.replace("\x00", "")

    cleaned_chars = []

    for char in decoded:
        category = unicodedata.category(char)

        if char in "\r\n\t":
            cleaned_chars.append(char)
        elif category.startswith("C"):
            cleaned_chars.append("\n")
        else:
            cleaned_chars.append(char)

    cleaned = "".join(cleaned_chars)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

    start_candidates = [
        cleaned.find("BỘ "),
        cleaned.find("TỔNG "),
        cleaned.find("CỘNG HÒA"),
        cleaned.find("Số:"),
    ]

    starts = [pos for pos in start_candidates if pos >= 0]

    if starts:
        cleaned = cleaned[min(starts):]

    lines = [line.strip() for line in cleaned.splitlines()]
    lines = [line for line in lines if _is_useful_doc_line(line)]

    return ExtractedDocument(
        text="\n".join(lines),
        extractor="legacy-doc-utf16le",
        page_count=None,
        warning="Legacy .doc file extracted with best-effort UTF-16 fallback.",
    )


def load_local_document(
    file_path: Path,
    options: LoaderOptions,
) -> ExtractedDocument:
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path, options)

    if suffix == ".docx":
        return extract_text_from_docx(file_path)

    if suffix == ".doc":
        return extract_text_from_legacy_doc(file_path)

    if suffix in SUPPORTED_IMAGE_SUFFIXES:
        return extract_text_from_image(file_path, options)

    raise ValueError(f"Unsupported file format: {suffix}")


def fetch_url(
    url: str,
    timeout_seconds: int = DEFAULT_TIMEOUT_SECONDS,
    user_agent: str = DEFAULT_USER_AGENT,
) -> WebFetchedContent:
    """Download one URL with a size limit. This is a single-page loader, not a crawler."""
    headers = {
        "User-Agent": user_agent,
        "Accept": (
            "text/html,application/xhtml+xml,application/pdf,image/*,"
            "application/msword,"
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document,"
            "*/*;q=0.8"
        ),
    }

    with requests.get(
        url,
        headers=headers,
        timeout=timeout_seconds,
        allow_redirects=True,
        stream=True,
    ) as response:
        response.raise_for_status()

        content_type = response.headers.get("content-type", "")
        content_length = response.headers.get("content-length")

        if content_length and int(content_length) > MAX_DOWNLOAD_BYTES:
            raise ValueError(
                f"Downloaded file is too large: {content_length} bytes "
                f"(limit: {MAX_DOWNLOAD_BYTES} bytes)"
            )

        chunks: list[bytes] = []
        total_size = 0

        for chunk in response.iter_content(chunk_size=1024 * 1024):
            if not chunk:
                continue

            total_size += len(chunk)

            if total_size > MAX_DOWNLOAD_BYTES:
                raise ValueError(
                    f"Downloaded file is too large: exceeded {MAX_DOWNLOAD_BYTES} bytes"
                )

            chunks.append(chunk)

        return WebFetchedContent(
            body=b"".join(chunks),
            content_type=content_type,
            final_url=response.url,
            encoding=response.encoding,
        )


def guess_suffix_from_url_or_content_type(url: str, content_type: str) -> str:
    path = urlparse(url).path.lower()

    for suffix in (
        ".pdf",
        ".docx",
        ".doc",
        ".html",
        ".htm",
        ".png",
        ".jpg",
        ".jpeg",
        ".tif",
        ".tiff",
        ".bmp",
        ".webp",
    ):
        if path.endswith(suffix):
            return ".html" if suffix == ".htm" else suffix

    normalized_content_type = content_type.split(";", 1)[0].strip().lower()

    content_type_to_suffix = {
        "application/pdf": ".pdf",
        "application/msword": ".doc",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "text/html": ".html",
        "application/xhtml+xml": ".html",
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/tiff": ".tiff",
        "image/bmp": ".bmp",
        "image/webp": ".webp",
    }

    return content_type_to_suffix.get(normalized_content_type, ".html")


def extract_charset_from_content_type(content_type: str) -> str | None:
    match = re.search(r"charset=([\w.\-]+)", content_type, flags=re.IGNORECASE)
    return match.group(1) if match else None


def decode_html_bytes(
    body: bytes,
    content_type: str,
    response_encoding: str | None = None,
) -> str:
    encoding_candidates = [
        extract_charset_from_content_type(content_type),
        response_encoding,
        "utf-8",
        "cp1258",
        "latin1",
    ]

    for encoding in encoding_candidates:
        if not encoding:
            continue

        try:
            return body.decode(encoding)
        except UnicodeDecodeError:
            continue

    return body.decode("utf-8", errors="replace")


def clean_html_visible_text(soup: BeautifulSoup) -> str:
    for element in soup(["script", "style", "noscript", "svg", "canvas", "iframe"]):
        element.decompose()

    for comment in soup.find_all(string=lambda text: isinstance(text, Comment)):
        comment.extract()

    content_root = (
        soup.find("main")
        or soup.find("article")
        or soup.find(attrs={"role": "main"})
        or soup.body
        or soup
    )

    title = soup.title.get_text(" ", strip=True) if soup.title else ""
    text = content_root.get_text("\n", strip=True)

    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]

    deduplicated_lines: list[str] = []
    previous = ""

    for line in lines:
        if line == previous:
            continue

        deduplicated_lines.append(line)
        previous = line

    if title and (not deduplicated_lines or deduplicated_lines[0] != title):
        deduplicated_lines.insert(0, title)

    return "\n".join(deduplicated_lines).strip()


def extract_text_from_html(
    body: bytes,
    content_type: str,
    final_url: str,
    response_encoding: str | None = None,
) -> ExtractedDocument:
    html = decode_html_bytes(
        body=body,
        content_type=content_type,
        response_encoding=response_encoding,
    )

    soup = BeautifulSoup(html, "html.parser")
    text = clean_html_visible_text(soup)

    return ExtractedDocument(
        text=text,
        extractor="web-html",
        page_count=None,
        warning=f"Fetched HTML from URL: {final_url}",
    )


def extract_text_from_url(
    url: str,
    options: LoaderOptions,
) -> ExtractedDocument:
    fetched = fetch_url(
        url=url,
        timeout_seconds=options.timeout_seconds,
        user_agent=options.user_agent,
    )

    suffix = guess_suffix_from_url_or_content_type(
        url=fetched.final_url,
        content_type=fetched.content_type,
    )

    if suffix == ".html":
        return extract_text_from_html(
            body=fetched.body,
            content_type=fetched.content_type,
            final_url=fetched.final_url,
            response_encoding=fetched.encoding,
        )

    supported_suffixes = {".pdf", ".docx", ".doc"} | SUPPORTED_IMAGE_SUFFIXES

    if suffix not in supported_suffixes:
        raise ValueError(
            f"Unsupported URL content type: {fetched.content_type or 'unknown'} "
            f"for URL: {fetched.final_url}"
        )

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_file = Path(temp_dir) / f"downloaded_document{suffix}"
        temp_file.write_bytes(fetched.body)
        extracted = load_local_document(temp_file, options)

    warning_parts = [
        f"Fetched {suffix.upper().lstrip('.')} from URL: {fetched.final_url}"
    ]

    if extracted.warning:
        warning_parts.append(extracted.warning)

    return ExtractedDocument(
        text=extracted.text,
        extractor=f"web-{extracted.extractor}",
        page_count=extracted.page_count,
        warning=join_warnings(*warning_parts),
    )


def load_document(
    source: Path | str,
    options: LoaderOptions,
) -> ExtractedDocument:
    source_value = str(source).strip()

    if is_url(source_value):
        return extract_text_from_url(source_value, options)

    return load_local_document(Path(source_value), options)


def safe_filename(document_id: str) -> str:
    cleaned = document_id.strip().replace("/", "_").replace("\\", "_")
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", cleaned).strip("_")


def actual_text_length(text: str) -> int:
    return meaningful_text_length(text)


def read_registry(registry_path: Path) -> pd.DataFrame:
    df = pd.read_excel(registry_path, dtype=str).fillna("")

    required_columns = ["document_id"]
    missing = [column for column in required_columns if column not in df.columns]

    if missing:
        raise ValueError(
            "Missing required column(s) in document_registry.xlsx: "
            + ", ".join(missing)
        )

    if "local_path" not in df.columns and "source_url" not in df.columns:
        raise ValueError(
            "document_registry.xlsx must contain either local_path or source_url."
        )

    return df


def choose_source_reference(row: pd.Series) -> str:
    local_path = str(row.get("local_path", "")).strip()
    source_url = str(row.get("source_url", "")).strip()

    if local_path:
        return local_path

    if is_url(source_url):
        return source_url

    return ""


def resolve_local_source_path(source_reference: str) -> Path:
    source_path = Path(source_reference)

    if source_path.is_absolute():
        return source_path

    return PROJECT_ROOT / source_path


def write_log(log_path: Path, rows: list[dict[str, object]]) -> None:
    log_path.parent.mkdir(parents=True, exist_ok=True)

    fieldnames = [
        "document_id",
        "local_path",
        "output_path",
        "status",
        "extractor",
        "page_count",
        "char_count",
        "warning",
        "error",
    ]

    with log_path.open("w", encoding="utf-8-sig", newline="") as file:
        writer = csv.DictWriter(file, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def run_loader(
    registry_path: Path,
    output_dir: Path,
    log_path: Path,
    skip_existing: bool = False,
    options: LoaderOptions | None = None,
) -> int:
    options = options or LoaderOptions()

    output_dir.mkdir(parents=True, exist_ok=True)
    df = read_registry(registry_path)

    logs: list[dict[str, object]] = []
    counters = {"success": 0, "empty": 0, "error": 0, "skipped": 0}

    for row_index, row in df.iterrows():
        document_id = str(row["document_id"]).strip()
        source_reference = choose_source_reference(row)

        if not document_id or not source_reference:
            print(
                f"[SKIP] Row {row_index + 2}: missing document_id or source "
                "(local_path/source_url)"
            )
            counters["skipped"] += 1
            continue

        output_path = output_dir / f"{safe_filename(document_id)}.txt"

        if skip_existing and output_path.exists():
            print(f"[SKIP] Existing output: {output_path}")
            counters["skipped"] += 1
            continue

        log_row: dict[str, object] = {
            "document_id": document_id,
            "local_path": source_reference,
            "output_path": str(output_path.relative_to(PROJECT_ROOT)),
            "status": "",
            "extractor": "",
            "page_count": "",
            "char_count": 0,
            "warning": "",
            "error": "",
        }

        try:
            print(f"[LOAD] {document_id} <- {source_reference}")

            if is_url(source_reference):
                extracted = load_document(source_reference, options)
            else:
                source_path = resolve_local_source_path(source_reference)

                if not source_path.exists():
                    raise FileNotFoundError(f"Source file not found: {source_path}")

                extracted = load_document(source_path, options)

            text = extracted.text
            char_count = actual_text_length(text)

            output_path.write_text(text, encoding="utf-8")

            status = "success" if char_count > 0 else "empty"

            if status == "empty":
                print(f"[WARNING] Empty extracted text: {document_id}")
            else:
                print(f"[OK] {document_id}: {char_count} chars -> {output_path}")

            log_row.update(
                {
                    "status": status,
                    "extractor": extracted.extractor,
                    "page_count": extracted.page_count or "",
                    "char_count": char_count,
                    "warning": extracted.warning,
                }
            )

            logs.append(log_row)
            counters[status] += 1

        except Exception as exc:
            print(f"[ERROR] {document_id}: {exc}")
            log_row["status"] = "error"
            log_row["error"] = str(exc)
            logs.append(log_row)
            counters["error"] += 1

    write_log(log_path, logs)

    print("\n===== DOCUMENT LOADER SUMMARY =====")
    print(f"Success: {counters['success']}")
    print(f"Empty:   {counters['empty']}")
    print(f"Error:   {counters['error']}")
    print(f"Skipped: {counters['skipped']}")
    print(f"Log:     {log_path}")

    return 1 if counters["error"] else 0


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Extract raw text from local files or URLs listed in document_registry.xlsx. "
            "Supports PDF, DOC, DOCX, HTML, image OCR, and scanned PDF OCR."
        )
    )

    parser.add_argument("--registry", type=Path, default=REGISTRY_PATH)
    parser.add_argument("--output-dir", type=Path, default=OUTPUT_DIR)
    parser.add_argument("--log-path", type=Path, default=LOG_PATH)

    parser.add_argument(
        "--skip-existing",
        action="store_true",
        help="Do not overwrite existing extracted .txt files.",
    )

    parser.add_argument(
        "--timeout",
        type=int,
        default=DEFAULT_TIMEOUT_SECONDS,
        help="HTTP timeout in seconds for URL sources.",
    )

    parser.add_argument(
        "--user-agent",
        type=str,
        default=DEFAULT_USER_AGENT,
        help="User-Agent header used when fetching URL sources.",
    )

    parser.add_argument(
        "--no-ocr",
        action="store_true",
        help="Disable OCR fallback for scanned PDFs and image files.",
    )

    parser.add_argument(
        "--force-ocr",
        action="store_true",
        help="Force OCR for every PDF even if normal text extraction succeeds.",
    )

    parser.add_argument(
        "--ocr-language",
        type=str,
        default=DEFAULT_OCR_LANGUAGE,
        help="Tesseract language code, e.g. vie, eng, or vie+eng.",
    )

    parser.add_argument(
        "--ocr-dpi",
        type=int,
        default=DEFAULT_OCR_DPI,
        help="DPI used when rendering PDF pages for OCR.",
    )

    parser.add_argument(
        "--ocr-min-chars",
        type=int,
        default=DEFAULT_OCR_MIN_CHARS,
        help=(
            "Run OCR automatically when normal PDF extraction returns fewer "
            "meaningful characters than this threshold."
        ),
    )

    return parser.parse_args()


def main() -> int:
    args = parse_args()

    options = LoaderOptions(
        timeout_seconds=args.timeout,
        user_agent=args.user_agent,
        ocr_enabled=not args.no_ocr,
        force_ocr=args.force_ocr,
        ocr_language=args.ocr_language,
        ocr_dpi=args.ocr_dpi,
        ocr_min_chars=args.ocr_min_chars,
    )

    return run_loader(
        registry_path=args.registry,
        output_dir=args.output_dir,
        log_path=args.log_path,
        skip_existing=args.skip_existing,
        options=options,
    )


if __name__ == "__main__":
    sys.exit(main())